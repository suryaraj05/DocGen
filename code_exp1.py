# document_generator.py

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def generate_document(input_text: str, output_path: str):
    doc = Document()
    
    # Set margins to 1.27 cm (approximately 36 pt)
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)

    # Store headings for Table of Contents
    toc_entries = []

    def add_breaker(doc):
        # Add a paragraph with a border for visual separation
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        border_run = p.add_run()
        border_run.add_text(' ')
        border_run.font.size = Pt(1)  # Make the text small so it's just a line
        p_format.space_after = Pt(0)  # Remove space after
        p_format.space_before = Pt(0)  # Remove space before
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Customize the border style
        p_format.border_bottom_color = RGBColor(166, 166, 166)  # Border color (light gray)
        p_format.border_bottom_width = Pt(2)  # Border width

    def add_heading(text, level='main'):
        para = doc.add_paragraph()
        run = para.add_run(text)
        if level == 'main':  # Main Heading
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(236, 159, 5)  # #EC9F05
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run.font.name = 'Jojoba'
            # Store the ToC entry
            toc_entries.append((text, level))
        elif level == 'sub':  # Sub Heading
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run.font.name = 'Jojoba'
            # Store the ToC entry
            toc_entries.append((text, level))
        elif level == 'subsub':  # Sub Sub Heading
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(245, 187, 0)  # #F5BB00
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run.font.name = 'Jojoba'
            # Store the ToC entry
            toc_entries.append((text, level))

    def add_content(text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(166, 166, 166)  # #A6A6A6
        run.font.name = 'Jojoba'
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_bullet(text, level='main'):
        para = doc.add_paragraph()
        if level == 'main':  # Main bullet
            bullet_text = f'• {text.strip()}'  # Use filled bullet
            para.add_run(bullet_text)
            para.paragraph_format.left_indent = Pt(12)  # Indent for main bullet
        elif level == 'sub':  # Sub bullet
            bullet_text = f'◦ {text.strip()}'  # Use outlined bullet
            para.add_run(bullet_text)
            para.paragraph_format.left_indent = Pt(24)  # Indent for sub bullet

        run = para.runs[0]
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(166, 166, 166)  # #A6A6A6
        run.font.name = 'Jojoba'

    def add_link(url):
        para = doc.add_paragraph()
        run = para.add_run(url)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 112, 192)  # #0070C0
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run.font.name = 'Jojoba'

    def add_toc(doc):
        # Add a Table of Contents
        doc.add_page_break()  # Add a page break before the TOC
        toc_heading = doc.add_heading('Table of Contents', level=1)
        toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for entry in toc_entries:
            title, level = entry
            # Create a new paragraph for the ToC entry
            toc_entry = doc.add_paragraph()
            toc_entry.add_run(title).bold = True
            if level == 'main':
                toc_entry.style = 'ListBullet'  # Main bullet style
            elif level == 'sub':
                toc_entry.style = 'ListBullet2'  # Sub bullet style
            elif level == 'subsub':
                toc_entry.style = 'ListBullet3'  # Sub sub bullet style

    # Process the input text
    lines = input_text.splitlines()

    # Process the input line by line
    for line in lines:
        line = line.strip()  # Remove leading/trailing whitespace

        # Extract the content inside # # using regex
        match = re.match(r'#\s*(.*?)\s*#', line)
        if match:
            content = match.group(1)

            # Handle Main Headings
            if content.startswith('Main Heading:'):
                heading_text = content.replace('Main Heading:', '').strip()
                add_heading(heading_text, level='main')
                add_breaker(doc)

            # Handle Sub Headings
            elif content.startswith('Sub Heading:'):
                sub_heading_text = content.replace('Sub Heading:', '').strip()
                add_heading(sub_heading_text, level='sub')
                add_breaker(doc)

            # Handle Sub Sub Headings
            elif content.startswith('Sub Sub Heading:'):
                sub_sub_heading_text = content.replace('Sub Sub Heading:', '').strip()
                add_heading(sub_sub_heading_text, level='subsub')

            # Handle Main Bullet Points (single dash `-`)
            elif content.startswith('- ') and content.endswith(' -'):  # Main bullet
                bullet_text = content[2:-2].strip()  # Remove leading "- " and trailing " -"
                add_bullet(bullet_text, level='main')

            # Handle Sub Bullet Points (double dash `--`)
            elif content.startswith('-- ') and content.endswith(' --'):  # Sub bullet
                bullet_text = content[3:-3].strip()  # Remove leading "-- " and trailing " --"
                add_bullet(bullet_text, level='sub')

            # Handle Links
            elif content.startswith('http'):  # Links
                add_link(content)

            # Handle Content (any text that is not a bullet or heading)
            else:
                add_content(content)

    # Add Table of Contents at the end
    add_toc(doc)

    # Save the document
    doc.save(output_path)
